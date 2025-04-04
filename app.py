import os
import re
import json
import time
import requests
from functools import wraps
from dotenv import load_dotenv
from flask import Flask, request, render_template, jsonify, send_file, Response, stream_with_context
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document
import logging
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT
from reportlab.lib.units import inch
import asyncio
import aiohttp
from typing import Dict, List
import hashlib
from asgiref.sync import async_to_sync
from concurrent.futures import ThreadPoolExecutor
from queue import Queue
from itertools import zip_longest
import openai

# 设置详细的日志格式
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - [%(filename)s:%(lineno)d] - %(message)s',
    handlers=[
        logging.FileHandler('app.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)

# 加载环境变量
load_dotenv()

# API配置
API_KEY = os.getenv("DEEPSEEK_API_KEY")
API_URL = "https://api.deepseek.com/v1/chat/completions"
MODEL_NAME = "deepseek-chat"
API_BASE = "https://api.deepseek.com/v1"
REQUEST_TIMEOUT = 60

# 创建session用于复用连接
session = requests.Session()
session.headers.update({
    "Authorization": f"Bearer {API_KEY}",
    "Content-Type": "application/json"
})

# 注册中文字体
try:
    # 尝试注册系统中的中文字体
    pdfmetrics.registerFont(TTFont('SimSun', 'C:/Windows/Fonts/simsun.ttc'))
except:
    logging.warning("无法加载 SimSun 字体，将使用默认字体")

# 翻译提示模板
DIRECT_TEMPLATE = """
# 角色
你是一位有20多年英汉翻译经验的资深翻译家，精通所有行业术语，且中英双语的表达方式都有着深刻理解。

# 任务
执行忠实的直译：
1. 充分理解原文所在行业和领域，并使用恰当的术语进行翻译
2. 尽可能保持原文的句子结构和词序，对于复杂和专业术语，可以维持英文不做翻译
3. 特别注意：
   - 理解上下文，提供最为准确的翻译
   - 完整保留原文信息
   - 维持关键术语翻译的一致性
   - 如果遇到@符号，说明它的上一行是标题，需要使用markdown一级标题格式(# 标题)进行翻译，并保留@符号

# 输出格式
- 逐段翻译
- 保留原文的段落分隔逻辑
- 重点关注原文作者含义的准确表达
"""

LITERARY_TEMPLATE = """
# 角色
你是一位具有20多年经验的资深出版社编辑，专门从事英文至中文翻译的校对与润色工作，对中文语言模式和表达惯例有深入理解，擅长讲对译者的初稿进行校对并进行更本地化的润色表达。

# 任务
将上一阶段直译的内容转化为意译后的，更本地化表达的，自然的中文：
1. 按照中文语言模式重构段落结构：
   - 调整词序和句法以符合中文习惯
   - 根据需要拆分或合并句子
   - 使用恰当的中文表达和习语
2. 在保持准确性的同时确保流畅性和可读性：
   - 明确隐含意义
   - 让复杂概念易于中文读者理解
3. 专注于用自然的中文传达准确含义
4. 注意：
   - 如果遇到@符号，说明它的上一行是标题，需要使用markdown一级标题格式(# 标题)进行翻译，并保留@符号

# 输出格式
- 返回读起来自然、本地化的中文文本
- 保留直译的完整含义
- 优先考虑清晰度和可理解性
"""

STYLIZED_TEMPLATE = """
# 角色
你是一位精通诺贝尔奖级的写作大师，擅长鲁迅、莫言、余华、张爱玲、余秋雨等中国当代最优秀作家的写作风格，同时也精通《纽约时报》、《经济学人》、《自然》等世界顶级媒体的写作风格。你擅长以下领域的写作，并且会根据原文的领域和风格选择你需要的写作风格：

1. 文学散文：张爱玲、余秋雨、林清玄 - 优雅含蓄、意境深远
2. 科技文章：《科学》《自然》《麻省理工科技评论》 - 严谨专业、通俗易懂
3. 财经报道：《经济学人》《华尔街日报》《金融时报》 - 深入浅出、数据详实
4. 人文社科：《纽约时报》《大西洋月刊》《纽约客》 - 思想深邃、文笔优美
5. 学术论文：钱钟书、季羡林、王国维 - 严谨典雅、论证缜密
6. 科普写作：卡尔·萨根、比尔·布莱森、理查德·费曼 - 生动有趣、深入浅出
7. 历史叙事：司马迁、吴晗、黄仁宇 - 史料翔实、叙事生动
8. 政论文章：梁启超、胡适、林语堂 - 思想深刻、文笔犀利
9. 商业写作：彼得·德鲁克、吉姆·柯林斯、杰克·韦尔奇 - 实用性强、见解独到
10. 哲学思辨：王阳明、牟宗三、冯友兰 - 思维缜密、论述严谨

# 任务
帮助出版社对上一阶段得到的英文原文和出色的翻译文稿进行进一步的润色和优化。你将根据对文本的理解选择最合适的写作风格，并对文章中的每个段落进行重构和润色，力求得到最佳的中文表达。做得好你将获得1000万美元奖金。

特别注意：
- 如果遇到@符号，说明它的上一行是标题，需要使用markdown一级标题格式(# 标题)进行翻译，并保留@符号

请直接返回翻译结果，不要做任何解释或说明。
"""

def retry_on_error(max_retries=3, delay=2):
    """重试装饰器，处理API调用失败的情况"""
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            retries = 0
            last_error = None
            
            while retries < max_retries:
                try:
                    return func(*args, **kwargs)
                except Exception as e:
                    retries += 1
                    last_error = e
                    logging.warning(f"第 {retries} 次调用失败: {str(e)}")
                    
                    if retries < max_retries:
                        sleep_time = delay * (2 ** (retries - 1))  # 指数退避
                        logging.info(f"等待 {sleep_time} 秒后重试...")
                        time.sleep(sleep_time)
                    
            logging.error(f"在 {max_retries} 次尝试后仍然失败: {str(last_error)}")
            raise last_error
            
        return wrapper
    return decorator

def split_text(text, min_words=1000, max_words=2000):
    """
    将文本分割成chunks，每个chunk包含1000-2000个单词。
    优先在段落换行处分割，如果超过1500个单词仍未找到换行，则在句子结尾处分割。
    对于markdown文件，遇到一级标题时也作为分割点。
    """
    chunks = []
    current_chunk = ""
    current_words = 0
    
    paragraphs = text.split('\n')
    
    for paragraph in paragraphs:
        # 检查是否是markdown一级标题
        if re.match(r'^#\s+.+', paragraph):
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = paragraph + '\n'
            current_words = len(paragraph.split())
            continue
            
        words = paragraph.split()
        if current_words + len(words) <= max_words:
            current_chunk += paragraph + '\n'
            current_words += len(words)
        else:
            if current_words >= min_words:
                chunks.append(current_chunk.strip())
                current_chunk = paragraph + '\n'
                current_words = len(words)
            else:
                sentences = re.split(r'(?<=[.!?])\s+', paragraph)
                for sentence in sentences:
                    sentence_words = sentence.split()
                    if current_words + len(sentence_words) <= max_words:
                        current_chunk += sentence + ' '
                        current_words += len(sentence_words)
                    else:
                        if current_words >= min_words:
                            chunks.append(current_chunk.strip())
                            current_chunk = sentence + ' '
                            current_words = len(sentence_words)
                        else:
                            current_chunk += sentence + ' '
                            chunks.append(current_chunk.strip())
                            current_chunk = ""
                            current_words = 0
                current_chunk += '\n'
    
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    return chunks

# 创建Flask应用
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'static/uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# 确保上传文件夹存在
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# 允许的文件类型
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'md'}
MIME_TYPES = {
    'pdf': 'application/pdf',
    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'txt': 'text/plain',
    'md': 'text/markdown'
}

def allowed_file(filename):
    """检查文件类型是否允许上传"""
    # 检查文件扩展名
    if '.' not in filename:
        return False
        
    ext = filename.rsplit('.', 1)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        return False
        
    # 检查文件大小（限制为10MB）
    if request.content_length and request.content_length > 10 * 1024 * 1024:
        return False
        
    return True

def extract_text_from_pdf(file_path):
    """从PDF文件中提取文本"""
    text = ""
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file_path):
    """从Word文档中提取文本"""
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def extract_text_from_file(file_path):
    """从文件中提取文本"""
    try:
        ext = file_path.rsplit('.', 1)[1].lower()
        text = None
        
        if ext == 'pdf':
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = '\n'.join(page.extract_text() for page in reader.pages)
        elif ext == 'docx':
            doc = Document(file_path)
            text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
        elif ext in ['txt', 'md']:
            # 尝试不同的编码方式读取文本文件
            for encoding in ['utf-8', 'gbk', 'latin-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                        break
                except UnicodeDecodeError:
                    continue
            if text is None:
                raise Exception("无法以支持的编码方式读取文件")
        else:
            raise Exception("不支持的文件类型")
            
        if not text:
            raise Exception("文件内容为空")
            
        return text
            
    except Exception as e:
        logging.error(f"提取文本时出错: {str(e)}")
        raise

@retry_on_error(max_retries=3, delay=2)
def call_deepseek_api(prompt, template):
    """调用 DeepSeek API 进行翻译"""
    try:
        logging.info(f"正在调用DeepSeek API... 模板长度: {len(template)}, 提示词长度: {len(prompt)}")
        
        start_time = time.time()
        
        # 准备请求数据
        payload = {
            "model": MODEL_NAME,
            "messages": [
                {"role": "system", "content": template},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.3,  # 降低随机性，使翻译更稳定
            "max_tokens": 4000,
            "stream": False
        }
        
        # 发送请求
        response = session.post(
            f"{API_BASE}/chat/completions",
            json=payload,
            timeout=REQUEST_TIMEOUT
        )
        
        # 检查响应状态
        response.raise_for_status()
        
        # 解析响应
        result = response.json()
        end_time = time.time()
        
        if "choices" in result and len(result["choices"]) > 0:
            content = result["choices"][0]["message"]["content"].strip()
            logging.info(f"API调用成功，耗时: {end_time - start_time:.2f}秒")
            return content
        else:
            raise Exception("API返回数据格式错误")
            
    except requests.exceptions.RequestException as e:
        logging.error(f"API请求失败: {str(e)}", exc_info=True)
        raise
    except Exception as e:
        logging.error(f"处理API响应时发生错误: {str(e)}", exc_info=True)
        raise

def process_text(text):
    """处理翻译结果，清理格式"""
    if not text:
        return text
        
    # 移除#和@符号
    lines = text.split('\n')
    processed_lines = []
    
    for line in lines:
        line = line.strip()
        if line:
            line = line.replace('#', '').replace('@', '').strip()
        processed_lines.append(line)
    
    return '\n'.join(processed_lines)

def translate_text_direct(text):
    """使用DeepSeek API进行直接翻译"""
    try:
        if isinstance(text, bytes):
            text = text.decode('utf-8')
        
        user_prompt = f"请翻译以下英文文本：\n{text}"
        
        result = call_deepseek_api(user_prompt, DIRECT_TEMPLATE)
        return process_text(result)
    except Exception as e:
        logging.error(f"直接翻译时发生错误: {str(e)}")
        raise

def translate_text_literary(text, first_translation):
    """使用DeepSeek API进行文学性翻译"""
    try:
        if isinstance(text, bytes):
            text = text.decode('utf-8')
        if isinstance(first_translation, bytes):
            first_translation = first_translation.decode('utf-8')
        
        user_prompt = f"""
        英文原文：
        {text}

        原翻译：
        {first_translation}

        只需返回翻译结果，不要做任何解释或说明。
        """
        
        result = call_deepseek_api(user_prompt, LITERARY_TEMPLATE)
        return process_text(result)
    except Exception as e:
        logging.error(f"文学性翻译时发生错误: {str(e)}")
        raise

def translate_text_stylized(text, literary_translation):
    """使用DeepSeek API进行风格化翻译"""
    try:
        if isinstance(text, bytes):
            text = text.decode('utf-8')
        if isinstance(literary_translation, bytes):
            literary_translation = literary_translation.decode('utf-8')
        
        user_prompt = f"""
        英文原文：
        {text}

        原翻译：
        {literary_translation}

        只需返回润色和风格化之后的翻译结果，不要做任何解释或说明。
        """
        
        result = call_deepseek_api(user_prompt, STYLIZED_TEMPLATE)
        return process_text(result)
    except Exception as e:
        logging.error(f"风格化翻译时发生错误: {str(e)}")
        raise

class DocumentFormatter:
    """文档格式化处理类，用于在翻译和PDF生成之间提供中间层处理"""
    
    def __init__(self):
        self.sections = []
        self.current_section = None
        self.last_content_type = None  # 记录上一个内容的类型
    
    def process_content(self, original, translation):
        """处理原文和翻译内容，将其组织成结构化的章节"""
        # 预处理：移除多余空行并规范化空白
        original_lines = self._preprocess_text(original)
        translation_lines = self._preprocess_text(translation)
        
        # 重置处理状态
        self.sections = []
        self.current_section = None
        self.last_content_type = None
        
        # 配对处理原文和翻译
        for orig, trans in zip_longest(original_lines, translation_lines, fillvalue=''):
            # 处理标题
            if orig.startswith('# ') or (orig and orig.endswith('@')):
                self._handle_title(orig, trans)
            else:
                # 只有当文本非空时才处理
                if orig.strip() or trans.strip():
                    self._handle_content(orig, trans)
                
        # 确保最后一个section被添加
        if self.current_section and self.current_section['content']:
            self.sections.append(self.current_section)
            
        return self.sections
    
    def _preprocess_text(self, text):
        """预处理文本，规范化空行和空白"""
        # 1. 分割成行
        lines = text.split('\n')
        
        # 2. 清理每行的首尾空白
        lines = [line.strip() for line in lines]
        
        # 3. 移除连续的空行，但保留段落之间的单个空行
        processed_lines = []
        prev_empty = False
        for line in lines:
            if line:
                processed_lines.append(line)
                prev_empty = False
            elif not prev_empty and processed_lines:  # 只在有内容后保留空行
                processed_lines.append(line)
                prev_empty = True
        
        # 4. 移除开头和结尾的空行
        while processed_lines and not processed_lines[0]:
            processed_lines.pop(0)
        while processed_lines and not processed_lines[-1]:
            processed_lines.pop()
            
        return processed_lines
    
    def _handle_title(self, orig, trans):
        """处理标题内容"""
        # 保存当前section（如果存在且有内容）
        if self.current_section and self.current_section['content']:
            self.sections.append(self.current_section)
        
        # 清理标题格式
        orig = orig.replace('# ', '').replace('@', '').strip()
        trans = trans.replace('# ', '').replace('@', '').strip()
        
        # 创建新的section
        self.current_section = {
            'type': 'title',
            'content': [(orig, trans)]
        }
        self.last_content_type = 'title'
    
    def _handle_content(self, orig, trans):
        """处理正文内容"""
        if not self.current_section:
            self.current_section = {
                'type': 'content',
                'content': []
            }
        elif self.current_section['type'] == 'title':
            # 如果上一个是标题section，创建新的内容section
            self.sections.append(self.current_section)
            self.current_section = {
                'type': 'content',
                'content': []
            }
        
        # 只添加非空的内容
        if orig.strip() or trans.strip():
            # 检查是否需要添加额外的空行
            if self.last_content_type == 'content' and self.current_section['content']:
                # 如果上一个内容也是正文，且当前section不为空，添加一个空行
                self.current_section['content'].append(('', ''))
            
            self.current_section['content'].append((orig.strip(), trans.strip()))
            self.last_content_type = 'content'

def create_pdf(translations, output_file):
    """生成PDF文件，使用DocumentFormatter处理格式"""
    try:
        doc = SimpleDocTemplate(
            output_file,
            pagesize=A4,
            rightMargin=50,
            leftMargin=50,
            topMargin=50,
            bottomMargin=50
        )

        # 创建样式
        styles = getSampleStyleSheet()
        
        # 标题样式
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontName='SimSun' if 'SimSun' in pdfmetrics.getRegisteredFontNames() else 'Helvetica',
            fontSize=14,
            spaceAfter=2,      # 进一步减少标题后的空间
            spaceBefore=12,    # 保持标题前的空间
            textColor=colors.HexColor('#2c3e50'),
            alignment=TA_LEFT,
            leading=16,
            keepWithNext=True  # 确保标题与下文在同一页
        )
        
        # 英文样式
        english_style = ParagraphStyle(
            'English',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            spaceBefore=0,     # 移除段落前空间
            spaceAfter=1,      # 最小化段落后空间
            textColor=colors.HexColor('#2c3e50'),
            alignment=TA_JUSTIFY,
            keepWithNext=True  # 确保英文与中文翻译在同一页
        )
        
        # 中文样式
        chinese_style = ParagraphStyle(
            'Chinese',
            parent=styles['Normal'],
            fontName='SimSun' if 'SimSun' in pdfmetrics.getRegisteredFontNames() else 'Helvetica',
            fontSize=10,
            leading=14,
            firstLineIndent=0,  # 移除首行缩进
            leftIndent=0,       # 移除左侧缩进
            spaceBefore=0,      # 移除段落前空间
            spaceAfter=2,       # 最小化段落后空间
            textColor=colors.HexColor('#34495e'),
            alignment=TA_JUSTIFY
        )

        # 构建文档内容
        story = []
        formatter = DocumentFormatter()
        
        # 处理每个翻译部分
        for trans in translations:
            sections = formatter.process_content(trans['original'], trans['stylized'])
            
            for i, section in enumerate(sections):
                if section['type'] == 'title':
                    # 如果不是第一个标题，添加额外的空间
                    if i > 0:
                        story.append(Spacer(1, 8))
                    
                    # 添加标题
                    orig, trans_text = section['content'][0]
                    if orig:
                        story.append(Paragraph(orig, title_style))
                    if trans_text:
                        story.append(Paragraph(trans_text, title_style))
                else:
                    # 添加正文内容
                    for orig, trans_text in section['content']:
                        if orig:
                            # 处理英文段落，移除多余空格
                            orig = ' '.join(orig.split())
                            story.append(Paragraph(orig, english_style))
                        if trans_text:
                            # 处理中文段落，移除多余空格和缩进
                            trans_text = ''.join(trans_text.split())
                            story.append(Paragraph(trans_text, chinese_style))
                
                # 只在内容section之后添加额外的空间
                if section['type'] == 'content':
                    story.append(Spacer(1, 2))

        # 生成PDF
        doc.build(story)
        return True
        
    except Exception as e:
        logging.error(f"生成PDF时发生错误: {str(e)}", exc_info=True)
        raise

class OptimizedTranslator:
    def __init__(self):
        self.session = requests.Session()
        self.session.headers.update({
            'Authorization': f'Bearer {API_KEY}',
            'Content-Type': 'application/json'
        })
        self.cache = {}
        self.processed_chunks = 0
        self.total_chunks = 0
        
    def _get_cache_key(self, text, translation_type):
        return hashlib.md5(f"{text}:{translation_type}".encode()).hexdigest()

    def _call_api(self, text, translation_type):
        """调用API进行翻译，带有重试机制"""
        cache_key = self._get_cache_key(text, translation_type)
        if cache_key in self.cache:
            return self.cache[cache_key]
            
        for attempt in range(3):
            try:
                start_time = time.time()
                
                # 根据翻译类型选择对应的模板
                if translation_type == 'direct':
                    template = DIRECT_TEMPLATE
                elif translation_type == 'literary':
                    template = LITERARY_TEMPLATE
                else:  # stylized
                    template = STYLIZED_TEMPLATE
                
                response = self.session.post(
                    f"{API_BASE}/chat/completions",
                    json={
                        'model': MODEL_NAME,
                        'messages': [
                            {'role': 'system', 'content': template},
                            {'role': 'user', 'content': text}
                        ],
                        'temperature': 0.3,
                        'max_tokens': 4000,
                        'stream': False
                    },
                    timeout=REQUEST_TIMEOUT
                )
                response.raise_for_status()
                result = response.json()['choices'][0]['message']['content']
                
                logging.info(f"API调用成功，耗时: {time.time() - start_time:.2f}秒")
                self.cache[cache_key] = result
                return result
                
            except Exception as e:
                if attempt == 2:  # 最后一次尝试失败
                    logging.error(f"API调用失败: {str(e)}")
                    raise
                logging.warning(f"第 {attempt + 1} 次调用失败: {str(e)}")
                time.sleep(2 ** attempt)

    def translate_chunk(self, text):
        """翻译一个文本块"""
        try:
            # 按顺序执行三种翻译
            direct = self._call_api(text, 'direct')
            literary = self._call_api(f"英文原文：\n{text}\n\n原翻译：\n{direct}\n\n只需返回翻译结果，不要做任何解释或说明。", 'literary')
            stylized = self._call_api(f"英文原文：\n{text}\n\n原翻译：\n{literary}\n\n只需返回润色和风格化之后的翻译结果，不要做任何解释或说明。", 'stylized')
            
            return {
                'original': text,
                'direct': direct,
                'literary': literary,
                'stylized': stylized
            }
            
        except Exception as e:
            logging.error(f"翻译文本块时出错: {str(e)}")
            raise

    def translate_text(self, text):
        """翻译完整文本，返回生成器以支持进度更新"""
        try:
            # 分割文本
            chunks = split_text(text)
            self.total_chunks = len(chunks)
            self.processed_chunks = 0
            
            results = []
            # 处理每个文本块
            for chunk in chunks:
                try:
                    result = self.translate_chunk(chunk)
                    results.append(result)
                    self.processed_chunks += 1
                    
                    # 发送进度更新
                    yield {
                        'type': 'progress',
                        'progress': self.processed_chunks / self.total_chunks
                    }
                    
                except Exception as e:
                    logging.error(f"处理翻译任务时出错: {str(e)}")
                    yield {
                        'type': 'error',
                        'error': str(e)
                    }
                    return
            
            # 发送完成消息
            yield {
                'type': 'completed',
                'translations': results
            }
            
        except Exception as e:
            logging.error(f"翻译过程中发生错误: {str(e)}")
            yield {
                'type': 'error',
                'error': str(e)
            }
            
    def __del__(self):
        """清理资源"""
        self.session.close()

class BatchTranslator:
    """批量翻译处理器，用于优化翻译速度和格式"""
    
    def __init__(self, translator, batch_size=3):
        self.translator = translator
        self.batch_size = batch_size
        self.cache = {}  # 用于存储已翻译的内容
        
    def _merge_paragraphs(self, paragraphs):
        """合并段落，保持格式标记"""
        merged = []
        current_batch = []
        
        for para in paragraphs:
            # 如果是标题，单独处理
            if para.strip().startswith('#'):
                if current_batch:
                    merged.append('\n'.join(current_batch))
                    current_batch = []
                merged.append(para)
                continue
                
            current_batch.append(para)
            
            # 当达到批量大小时合并
            if len(current_batch) >= self.batch_size:
                merged.append('\n'.join(current_batch))
                current_batch = []
        
        # 处理剩余的段落
        if current_batch:
            merged.append('\n'.join(current_batch))
            
        return merged
    
    def _split_translation(self, merged_text, original_paragraphs):
        """将合并后的翻译结果分割回原段落"""
        # 使用原文作为分隔符来分割翻译结果
        result = []
        current_pos = 0
        
        for para in original_paragraphs:
            if para.strip().startswith('#'):
                # 标题直接添加
                result.append(para)
                continue
                
            # 在合并文本中查找当前段落
            para_start = merged_text.find(para, current_pos)
            if para_start != -1:
                # 找到下一段的开始位置
                next_para_start = merged_text.find('\n\n', para_start)
                if next_para_start == -1:
                    next_para_start = len(merged_text)
                    
                # 提取翻译后的段落
                translated_para = merged_text[para_start:next_para_start].strip()
                result.append(translated_para)
                current_pos = next_para_start
            else:
                # 如果找不到原文，保持原样
                result.append(para)
                
        return result
    
    def translate_batch(self, text):
        """批量翻译文本"""
        # 分割成段落
        paragraphs = text.split('\n\n')
        paragraphs = [p.strip() for p in paragraphs if p.strip()]
        
        # 合并段落
        merged_paragraphs = self._merge_paragraphs(paragraphs)
        
        # 批量翻译
        translated_paragraphs = []
        for batch in merged_paragraphs:
            # 检查缓存
            if batch in self.cache:
                translated_paragraphs.append(self.cache[batch])
                continue
                
            # 翻译并缓存
            translated = self.translator.translate_text(batch)
            self.cache[batch] = translated
            translated_paragraphs.append(translated)
        
        # 合并翻译结果
        merged_translation = '\n\n'.join(translated_paragraphs)
        
        # 分割回原段落
        final_translation = self._split_translation(merged_translation, paragraphs)
        
        return '\n\n'.join(final_translation)

def translate_text(text, target_lang='zh'):
    """翻译文本"""
    try:
        # 使用批量翻译器
        batch_translator = BatchTranslator(translator)
        return batch_translator.translate_batch(text)
    except Exception as e:
        logger.error(f"翻译失败: {str(e)}")
        raise

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    """处理文件上传和翻译请求"""
    temp_file_path = None
    file = None
    try:
        # 检查是否有文件被上传
        if 'file' not in request.files:
            return jsonify({'error': '没有文件被上传'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '没有选择文件'}), 400
            
        # 检查文件类型
        if not allowed_file(file.filename):
            return jsonify({'error': '不支持的文件类型'}), 400
            
        # 安全地获取文件名
        filename = secure_filename(file.filename)
        if not filename:
            return jsonify({'error': '无效的文件名'}), 400
            
        # 确保上传目录存在
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        
        # 保存文件
        temp_file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(temp_file_path)
        file.close()  # 显式关闭文件
        
        # 提取文本
        try:
            text = extract_text_from_file(temp_file_path)
            if not text:
                return jsonify({'error': '无法从文件中提取文本'}), 400
            logging.info(f"成功提取文本，长度: {len(text)} 字符")
        except Exception as e:
            logging.error(f"提取文本时出错: {str(e)}")
            return jsonify({'error': '提取文本失败'}), 500
            
        def generate():
            output_file_path = None
            try:
                # 创建翻译器实例
                translator = OptimizedTranslator()
                
                # 处理翻译结果
                for update in translator.translate_text(text):
                    if update['type'] == 'completed':
                        try:
                            # 生成PDF
                            translations = update['translations']
                            output_filename = f"translated_{filename.rsplit('.', 1)[0]}.pdf"
                            output_file_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
                            create_pdf(translations, output_file_path)
                            
                            # 发送完成消息
                            yield json.dumps({
                                'type': 'completed',
                                'download_url': f'/download/{output_filename}'
                            }) + '\n'
                        except Exception as e:
                            logging.error(f"生成PDF时出错: {str(e)}")
                            yield json.dumps({
                                'type': 'error',
                                'error': '生成PDF失败'
                            }) + '\n'
                    else:
                        yield json.dumps(update) + '\n'
                        
            except Exception as e:
                logging.error(f"处理过程中发生错误: {str(e)}")
                yield json.dumps({
                    'type': 'error',
                    'error': str(e)
                }) + '\n'
            finally:
                # 清理临时文件
                try:
                    if temp_file_path and os.path.exists(temp_file_path):
                        # 确保文件句柄被释放
                        import gc
                        gc.collect()  # 强制垃圾回收
                        time.sleep(0.1)  # 给系统一点时间完成文件操作
                        os.remove(temp_file_path)
                except Exception as e:
                    logging.error(f"清理临时文件时出错: {str(e)}")

        return Response(stream_with_context(generate()), mimetype='text/event-stream')
        
    except Exception as e:
        logging.error(f"处理上传文件时发生错误: {str(e)}")
        # 清理临时文件
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                if file:
                    file.close()  # 确保文件被关闭
                # 确保文件句柄被释放
                import gc
                gc.collect()  # 强制垃圾回收
                time.sleep(0.1)  # 给系统一点时间完成文件操作
                os.remove(temp_file_path)
            except Exception as cleanup_error:
                logging.error(f"清理临时文件时出错: {str(cleanup_error)}")
        return jsonify({'error': str(e)}), 500

@app.route('/download/<filename>')
def download_file(filename):
    return send_file(
        os.path.join(app.config['UPLOAD_FOLDER'], filename),
        as_attachment=True
    )

@app.route('/convert_to_pdf/<path:filename>')
def convert_txt_to_pdf(filename):
    """将指定的txt文件转换为PDF"""
    try:
        input_file = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(input_file):
            return jsonify({'error': '文件不存在'}), 404

        # 读取文本文件
        try:
            with open(input_file, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(input_file, 'r', encoding='gbk') as f:
                content = f.read()

        # 准备输出文件名
        output_filename = f"{os.path.splitext(filename)[0]}.pdf"
        output_file = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)

        # 将内容分割成段落
        paragraphs = content.split('\n')

        # 创建PDF文档
        doc = SimpleDocTemplate(
            output_file,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        # 创建样式
        styles = getSampleStyleSheet()
        
        # 创建自定义样式
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontName='SimSun' if 'SimSun' in pdfmetrics.getRegisteredFontNames() else 'Helvetica',
            fontSize=24,
            spaceAfter=30,
            alignment=TA_LEFT
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontName='SimSun' if 'SimSun' in pdfmetrics.getRegisteredFontNames() else 'Helvetica',
            fontSize=16,
            spaceAfter=20,
            alignment=TA_LEFT
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName='SimSun' if 'SimSun' in pdfmetrics.getRegisteredFontNames() else 'Helvetica',
            fontSize=12,
            leading=14,
            alignment=TA_JUSTIFY
        )

        # 构建文档内容
        story = []
        
        # 添加标题
        story.append(Paragraph("研究技艺", title_style))
        story.append(Spacer(1, 24))

        current_section = None
        
        for paragraph in paragraphs:
            if not paragraph.strip():
                continue
                
            # 检查是否是章节标题（以#开头）
            if paragraph.strip().startswith('#'):
                level = len(paragraph.split('#')[0]) + 1
                title_text = paragraph.strip('#').strip()
                if level == 1:
                    # 在新章节前添加分页符（除了第一个章节）
                    if current_section is not None:
                        story.append(PageBreak())
                    current_section = title_text
                    story.append(Paragraph(title_text, title_style))
                else:
                    story.append(Paragraph(title_text, heading_style))
                story.append(Spacer(1, 12))
            else:
                # 普通段落
                try:
                    story.append(Paragraph(paragraph, normal_style))
                    story.append(Spacer(1, 6))
                except Exception as e:
                    logging.warning(f"处理段落时出错: {str(e)}")
                    continue

        # 生成PDF
        doc.build(story)
        
        return send_file(
            output_file,
            as_attachment=True,
            download_name=output_filename,
            mimetype='application/pdf'
        )

    except Exception as e:
        logging.error(f"转换PDF时发生错误: {str(e)}", exc_info=True)
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True) 